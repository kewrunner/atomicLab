#!/usr/bin/env python3
"""
Local Documents MCP Server
A secure, single-file MCP server for accessing local documents via AI assistants.
"""

import os
import re
import sys
import logging
from pathlib import Path
from typing import Annotated, Literal, List, Dict, Any, Optional
from datetime import datetime

from mcp.server.fastmcp import FastMCP
from openpyxl import load_workbook
from docx import Document
from pypdf import PdfReader


# ============================================================================
# Configuration
# ============================================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

SERVER_NAME = "Local Documents"
HOME_DIR = Path.home().resolve()

MAX_FILE_SIZE_MB: float = 5.0
ALLOWED_EXTENSIONS: Optional[List[str]] = None
BLOCKED_EXTENSIONS: List[str] = [
    ".exe", ".sh", ".bat", ".cmd", ".ps1", ".dll"
]
HIDDEN_FILES_ENABLED: bool = False
MAX_DIRECTORY_ITEMS: int = 500


# ============================================================================
# Security Utilities
# ============================================================================

class SecurityError(Exception):
    """Custom exception for security-related access violations."""
    pass


def _secure_path(relative_path: str) -> Path:
    """Resolve a relative or absolute path inside the home directory."""

    clean_path = relative_path.strip()

    if not clean_path or clean_path == "~":
        target_path = HOME_DIR
    else:
        requested_path = Path(clean_path).expanduser()
        if requested_path.is_absolute():
            target_path = requested_path.resolve()
        else:
            target_path = (HOME_DIR / requested_path).resolve()

    try:
        target_path.relative_to(HOME_DIR)
    except ValueError:
        raise SecurityError("Access denied: Path traversal detected.")

    return target_path


def _check_file_allowed(file_path: Path) -> None:
    """Verify that a file is allowed to be accessed."""

    suffix = file_path.suffix.lower()

    if suffix in BLOCKED_EXTENSIONS:
        raise SecurityError(
            f"Access denied: '{suffix}' files are not permitted."
        )

    if ALLOWED_EXTENSIONS is not None and suffix not in ALLOWED_EXTENSIONS:
        raise SecurityError(
            f"Access denied: '{suffix}' files are not on the allowed list."
        )


# ============================================================================
# Initialize MCP Server
# ============================================================================

mcp = FastMCP(SERVER_NAME)


# ============================================================================
# MCP Tools
# ============================================================================

@mcp.tool()
def list_directory(
    sub_directory: Annotated[
        str,
        "Relative path from home directory (empty = home)"
    ] = ""
) -> List[Dict[str, Any]]:
    """List directory contents with metadata."""

    logger.info(f"Listing directory: {sub_directory or '~'}")

    try:
        target_path = _secure_path(sub_directory)

        if not target_path.exists():
            return [{"error": f"Directory not found: '{sub_directory}'"}]

        if not target_path.is_dir():
            return [{"error": f"Not a directory: '{sub_directory}'"}]

        items = []
        count = 0

        for entry in target_path.iterdir():
            if not HIDDEN_FILES_ENABLED and entry.name.startswith("."):
                continue

            if count >= MAX_DIRECTORY_ITEMS:
                items.append({
                    "name": "...",
                    "type": "truncated",
                    "note": f"Output limited to {MAX_DIRECTORY_ITEMS} items"
                })
                break

            stat_info = entry.stat()

            items.append({
                "name": entry.name,
                "type": "directory" if entry.is_dir() else "file",
                "size_bytes": stat_info.st_size if entry.is_file() else 0,
                "modified": datetime.fromtimestamp(
                    stat_info.st_mtime
                ).isoformat()
            })

            count += 1

        items.sort(
            key=lambda item: (
                item["type"] != "directory",
                item["name"].lower()
            )
        )

        return items

    except SecurityError as error:
        logger.warning(f"Security violation: {error}")
        return [{"error": str(error)}]

    except Exception as error:
        logger.error(f"Error listing directory: {error}")
        return [{"error": f"Failed to list directory: {error}"}]


@mcp.tool()
def read_file(
    file_path: Annotated[
        str,
        "Relative path to the file from home directory"
    ],
    max_lines: Annotated[
        int,
        "Maximum lines to read (0 = no limit)"
    ] = 0,
    start_line: Annotated[
        int,
        "Starting line number (1-indexed)"
    ] = 1
) -> Dict[str, Any]:
    """Read a text file with optional line range selection."""

    logger.info(f"Reading file: {file_path}")

    try:
        target_path = _secure_path(file_path)

        if not target_path.exists():
            return {"error": f"File not found: '{file_path}'"}

        if not target_path.is_file():
            return {"error": f"Not a file: '{file_path}'"}

        _check_file_allowed(target_path)

        file_size_mb = target_path.stat().st_size / (1024 * 1024)

        if file_size_mb > MAX_FILE_SIZE_MB:
            return {
                "error": f"File too large ({file_size_mb:.1f}MB)",
                "max_allowed_mb": MAX_FILE_SIZE_MB,
                "suggestion": "Use max_lines and start_line to read portions"
            }

        text = target_path.read_text(
            encoding="utf-8",
            errors="replace"
        )

        if max_lines > 0:
            lines = text.splitlines()
            start = max(0, start_line - 1)
            end = start + max_lines

            return {
                "file_path": str(target_path),
                "total_lines": len(lines),
                "returned_lines": len(lines[start:end]),
                "start_line": start_line,
                "content": "\n".join(lines[start:end]),
                "truncated": end < len(lines)
            }

        return {
            "file_path": str(target_path),
            "content": text,
            "size_bytes": len(text.encode("utf-8")),
            "line_count": text.count("\n") + 1
        }

    except SecurityError as error:
        logger.warning(
            f"Security violation reading {file_path}: {error}"
        )
        return {"error": str(error)}

    except Exception as error:
        logger.error(f"Error reading file: {error}")
        return {"error": f"Failed to read file: {error}"}


@mcp.tool()
def read_document(
    file_path: Annotated[
        str,
        "Relative or absolute path to an Excel, CSV, Word, PDF, or text file"
    ],
    sheet_name: Annotated[
        Optional[str],
        "Excel worksheet name; omit to read the first worksheet"
    ] = None,
    start_row: Annotated[
        int,
        "First spreadsheet row to return, 1-indexed"
    ] = 1,
    max_rows: Annotated[
        int,
        "Maximum spreadsheet rows to return"
    ] = 500,
) -> Dict[str, Any]:
    """Read common document formats into structured data for analysis."""

    try:
        target_path = _secure_path(file_path)

        if not target_path.exists() or not target_path.is_file():
            return {"error": f"File not found: '{file_path}'"}

        _check_file_allowed(target_path)

        if target_path.stat().st_size > MAX_FILE_SIZE_MB * 1024 * 1024:
            return {"error": f"File exceeds {MAX_FILE_SIZE_MB}MB limit"}

        suffix = target_path.suffix.lower()

        if suffix in {".xlsx", ".xlsm"}:
            workbook = load_workbook(
                target_path,
                read_only=True,
                data_only=True
            )
            selected_sheet = sheet_name or workbook.sheetnames[0]

            if selected_sheet not in workbook.sheetnames:
                return {
                    "error": "Worksheet not found",
                    "available_sheets": workbook.sheetnames
                }

            worksheet = workbook[selected_sheet]
            rows = worksheet.iter_rows(values_only=True)
            headers = [str(value) if value is not None else "" for value in next(rows, ())]
            records = []

            for row_number, row in enumerate(rows, start=2):
                if row_number < start_row:
                    continue
                if len(records) >= max_rows:
                    break
                records.append({
                    headers[index] if index < len(headers) and headers[index] else f"column_{index + 1}": value
                    for index, value in enumerate(row)
                })

            workbook.close()
            return {
                "file_path": str(target_path),
                "format": "spreadsheet",
                "sheet": selected_sheet,
                "available_sheets": workbook.sheetnames,
                "columns": headers,
                "row_count_returned": len(records),
                "rows": records
            }

        if suffix == ".csv":
            import csv

            with target_path.open(
                "r",
                encoding="utf-8-sig",
                errors="replace",
                newline=""
            ) as csv_file:
                reader = csv.DictReader(csv_file)
                records = []
                for row_number, row in enumerate(reader, start=2):
                    if row_number < start_row:
                        continue
                    if len(records) >= max_rows:
                        break
                    records.append(dict(row))

            return {
                "file_path": str(target_path),
                "format": "csv",
                "columns": list(records[0].keys()) if records else [],
                "row_count_returned": len(records),
                "rows": records
            }

        if suffix == ".docx":
            document = Document(target_path)
            paragraphs = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]
            tables = []

            for table in document.tables:
                tables.append([
                    [cell.text for cell in row.cells]
                    for row in table.rows
                ])

            return {
                "file_path": str(target_path),
                "format": "word",
                "paragraphs": paragraphs,
                "tables": tables
            }

        if suffix == ".pdf":
            reader = PdfReader(str(target_path))
            pages = [page.extract_text() or "" for page in reader.pages]

            return {
                "file_path": str(target_path),
                "format": "pdf",
                "page_count": len(pages),
                "pages": pages,
                "text": "\n\n".join(pages)
            }

        if suffix in {".txt", ".md", ".json", ".yaml", ".yml"}:
            return {
                "file_path": str(target_path),
                "format": "text",
                "text": target_path.read_text(
                    encoding="utf-8",
                    errors="replace"
                )
            }

        return {
            "error": f"Unsupported document format: '{suffix}'",
            "supported_formats": [
                ".xlsx", ".xlsm", ".csv", ".docx", ".pdf",
                ".txt", ".md", ".json", ".yaml", ".yml"
            ]
        }

    except SecurityError as error:
        return {"error": str(error)}
    except Exception as error:
        logger.exception("Document reading failed")
        return {"error": f"Failed to read document: {error}"}


@mcp.tool()
def search_files(
    query: Annotated[
        str,
        "Search term or filename pattern"
    ],
    search_type: Annotated[
        Literal["content", "name"],
        "Search file contents or filenames"
    ] = "name",
    sub_directory: Annotated[
        str,
        "Directory to search within"
    ] = "",
    file_pattern: Annotated[
        str,
        "File extension filter, such as .txt or .md"
    ] = ""
) -> Dict[str, Any]:
    """Search files by name or content."""

    logger.info(
        f"Searching '{query}' in "
        f"{sub_directory or '~'} ({search_type})"
    )

    try:
        target_path = _secure_path(sub_directory)

        if not target_path.exists():
            return {"error": f"Directory not found: '{sub_directory}'"}

        if not target_path.is_dir():
            return {"error": f"Not a directory: '{sub_directory}'"}

        results = []
        files_checked = 0
        max_search_files = 100

        if search_type == "name":
            pattern = re.compile(query.lower())
        else:
            pattern = re.compile(query, re.IGNORECASE)

        for entry in target_path.rglob("*"):
            if (
                not HIDDEN_FILES_ENABLED
                and entry.name.startswith(".")
            ):
                continue

            if not entry.is_file():
                continue

            if file_pattern and not entry.match(f"*{file_pattern}"):
                continue

            relative_path = str(entry.relative_to(target_path))

            if search_type == "name":
                if pattern.search(entry.name.lower()):
                    results.append({
                        "path": relative_path,
                        "match": entry.name,
                        "type": "name_match"
                    })

                    if len(results) >= 50:
                        break

            else:
                if files_checked >= max_search_files:
                    continue

                try:
                    _check_file_allowed(entry)

                    if (
                        entry.stat().st_size
                        > MAX_FILE_SIZE_MB * 1024 * 1024
                    ):
                        continue

                    content = entry.read_text(
                        encoding="utf-8",
                        errors="replace"
                    )

                    matches = pattern.findall(content)

                    if matches:
                        lower_content = content.lower()
                        lower_query = query.lower()
                        match_position = lower_content.find(lower_query)

                        preview_start = max(0, match_position - 50)
                        preview_end = match_position + len(query) + 50

                        results.append({
                            "path": relative_path,
                            "match_count": len(matches),
                            "preview": content[
                                preview_start:preview_end
                            ],
                            "type": "content_match"
                        })

                except Exception:
                    continue

                files_checked += 1

        return {
            "query": query,
            "search_type": search_type,
            "directory": str(target_path),
            "results_count": len(results),
            "results": results[:50]
        }

    except SecurityError as error:
        logger.warning(f"Security violation: {error}")
        return {"error": str(error)}

    except Exception as error:
        logger.error(f"Search error: {error}")
        return {"error": f"Search failed: {error}"}


@mcp.tool()
def get_file_metadata(
    file_path: Annotated[
        str,
        "Relative path to the file or folder"
    ]
) -> Dict[str, Any]:
    """Get detailed metadata for a file or directory."""

    logger.info(f"Getting metadata: {file_path}")

    try:
        target_path = _secure_path(file_path)

        if not target_path.exists():
            return {"error": "Path does not exist"}

        stat_info = target_path.stat()

        permissions = {
            "readable": os.access(target_path, os.R_OK),
            "writable": os.access(target_path, os.W_OK),
            "executable": os.access(target_path, os.X_OK)
        }

        return {
            "name": target_path.name,
            "type": (
                "file" if target_path.is_file()
                else "directory"
            ),
            "size_bytes": stat_info.st_size,
            "size_formatted": _format_size(stat_info.st_size),
            "absolute_path": str(target_path),
            "extension": (
                target_path.suffix if target_path.is_file()
                else None
            ),
            "created": datetime.fromtimestamp(
                stat_info.st_ctime
            ).isoformat(),
            "modified": datetime.fromtimestamp(
                stat_info.st_mtime
            ).isoformat(),
            "accessed": datetime.fromtimestamp(
                stat_info.st_atime
            ).isoformat(),
            "permissions": permissions,
            "is_hidden": target_path.name.startswith(".")
        }

    except SecurityError as error:
        return {"error": str(error)}

    except Exception as error:
        return {
            "error": f"Failed to read metadata: {error}"
        }


# ============================================================================
# Utility Functions
# ============================================================================

def _format_size(size_bytes: int) -> str:
    """Return a human-readable file size."""

    size = float(size_bytes)

    for unit in ["B", "KB", "MB", "GB"]:
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024

    return f"{size:.1f} TB"


# ============================================================================
# Entry Point
# ============================================================================

if __name__ == "__main__":
    banner_lines = [
        f"{SERVER_NAME} MCP Server",
        f"Accessing: {HOME_DIR}",
        f"Max file size: {MAX_FILE_SIZE_MB:.1f}MB",
    ]

    width = max(
        60,
        max(len(line) for line in banner_lines)
    )

    print(
        "╔" + "═" * (width + 2) + "╗",
        file=sys.stderr
    )

    for line in banner_lines:
        print(
            f"║ {line:<{width}} ║",
            file=sys.stderr
        )

    print(
        "╚" + "═" * (width + 2) + "╝",
        file=sys.stderr
    )

    mcp.run()
